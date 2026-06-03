import streamlit as st
import nltk
import spacy
nltk.download('stopwords')

import pandas as pd
import base64, random
import time, datetime
from pdfminer3.layout import LAParams, LTTextBox
from pdfminer3.pdfpage import PDFPage
from pdfminer3.pdfinterp import PDFResourceManager
from pdfminer3.pdfinterp import PDFPageInterpreter
from pdfminer3.converter import TextConverter
import io, random, re, json, subprocess, os, tempfile
from streamlit_tags import st_tags
from PIL import Image
import pymysql
from Courses import ds_course, web_course, android_course, ios_course, uiux_course, resume_videos, interview_videos
from chatbot import render_chatbot
import plotly.express as px
from yt_dlp import YoutubeDL

# ── Custom Resume Parser (replaces broken pyresparser) ───────────────────────

nlp = spacy.load('en_core_web_sm')

def extract_email(text):
    match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    return match.group(0) if match else ''

def extract_mobile(text):
    match = re.search(r'(\+?\d[\d\s\-]{8,}\d)', text)
    return match.group(0).strip() if match else ''

def extract_name(text):
    doc = nlp(text[:500])
    for ent in doc.ents:
        if ent.label_ == 'PERSON':
            return ent.text
    for line in text.splitlines():
        line = line.strip()
        if line and len(line.split()) <= 5:
            return line
    return ''

SKILLS_DB = [
    'python', 'java', 'c', 'c++', 'c#', 'javascript', 'typescript', 'ruby',
    'php', 'swift', 'kotlin', 'go', 'rust', 'scala', 'r', 'matlab',
    'html', 'css', 'react', 'angular', 'vue', 'node js', 'node.js', 'django',
    'flask', 'spring', 'laravel', 'wordpress', 'magento',
    'sql', 'mysql', 'postgresql', 'mongodb', 'sqlite', 'oracle', 'redis',
    'machine learning', 'deep learning', 'tensorflow', 'keras', 'pytorch',
    'scikit-learn', 'pandas', 'numpy', 'data science', 'data analysis',
    'nlp', 'computer vision', 'opencv',
    'android', 'android development', 'flutter', 'xml', 'kivy',
    'ios', 'ios development', 'xcode', 'cocoa', 'objective-c',
    'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git',
    'linux', 'bash', 'powershell',
    'figma', 'adobe xd', 'photoshop', 'illustrator', 'after effects',
    'ux', 'ui', 'prototyping', 'wireframes', 'zeplin', 'balsamiq',
    'streamlit', 'power bi', 'tableau', 'excel', 'spark', 'hadoop',
]

def extract_skills(text):
    text_lower = text.lower()
    found = []
    for skill in SKILLS_DB:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            found.append(skill.title())
    return list(dict.fromkeys(found))

def count_pdf_pages(filepath):
    count = 0
    with open(filepath, 'rb') as f:
        for _ in PDFPage.get_pages(f, caching=True, check_extractable=True):
            count += 1
    return count if count else 1

def parse_resume(filepath):
    text = pdf_reader(filepath)
    return {
        'name':          extract_name(text),
        'email':         extract_email(text),
        'mobile_number': extract_mobile(text),
        'skills':        extract_skills(text),
        'no_of_pages':   count_pdf_pages(filepath),
    }

# ── Resume Generator (pure Python via python-docx) ───────────────────────────

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io as _io

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GRAY   = RGBColor(0x55, 0x55, 0x55)
BLACK  = RGBColor(0x00, 0x00, 0x00)

def _set_para_border_bottom(para, color="1F4E79", size=8):
    """Add a bottom border to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def _section_header(doc, title):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)
    _set_para_border_bottom(para)
    run = para.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT
    run.font.name = 'Arial'
    return para

def _bullet(doc, text):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.font.color.rgb = BLACK
    return para

def _date_row(doc, left, right):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    r1 = para.add_run(left)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.name = 'Arial'
    r2 = para.add_run('\t' + right)
    r2.font.size = Pt(9)
    r2.font.name = 'Arial'
    r2.font.color.rgb = GRAY
    para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    return para

def generate_resume_docx(d):
    """Generate a .docx resume from a dict and return the file bytes."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── Name ──────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(d.get('name') or 'Your Name')
    name_run.bold = True
    name_run.font.size = Pt(26)
    name_run.font.color.rgb = ACCENT
    name_run.font.name = 'Arial'

    # ── Contact line ───────────────────────────────────────────────────
    contact_parts = [d.get('email'), d.get('phone'), d.get('linkedin'), d.get('location')]
    contact = '  |  '.join(p for p in contact_parts if p)
    if contact:
        c_para = doc.add_paragraph()
        c_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_run = c_para.add_run(contact)
        c_run.font.size = Pt(9)
        c_run.font.name = 'Arial'
        c_run.font.color.rgb = GRAY

    # Divider
    div = doc.add_paragraph()
    div.paragraph_format.space_after = Pt(4)
    _set_para_border_bottom(div, size=12)

    # ── Summary ────────────────────────────────────────────────────────
    if d.get('summary'):
        _section_header(doc, 'Professional Summary')
        p = doc.add_paragraph()
        r = p.add_run(d['summary'])
        r.font.size = Pt(10)
        r.font.name = 'Arial'

    # ── Skills ─────────────────────────────────────────────────────────
    if d.get('skills'):
        _section_header(doc, 'Skills')
        p = doc.add_paragraph()
        r = p.add_run('  \u2022  '.join(d['skills']))
        r.font.size = Pt(10)
        r.font.name = 'Arial'

    # ── Experience ─────────────────────────────────────────────────────
    if d.get('experience'):
        _section_header(doc, 'Work Experience')
        for exp in d['experience']:
            left  = f"{exp.get('title','')}  \u2014  {exp.get('company','')}"
            right = f"{exp.get('start','')} \u2013 {exp.get('end','Present')}"
            _date_row(doc, left, right)
            if exp.get('location'):
                lp = doc.add_paragraph()
                lr = lp.add_run(exp['location'])
                lr.italic = True
                lr.font.size = Pt(9)
                lr.font.color.rgb = GRAY
                lr.font.name = 'Arial'
            for b in exp.get('bullets', []):
                _bullet(doc, b)
            doc.add_paragraph()

    # ── Education ──────────────────────────────────────────────────────
    if d.get('education'):
        _section_header(doc, 'Education')
        for edu in d['education']:
            left  = f"{edu.get('degree','')}  \u2014  {edu.get('institution','')}"
            right = f"{edu.get('start','')} \u2013 {edu.get('end','')}"
            _date_row(doc, left, right)
            if edu.get('gpa'):
                gp = doc.add_paragraph()
                gr = gp.add_run(f"GPA: {edu['gpa']}")
                gr.font.size = Pt(9)
                gr.font.color.rgb = GRAY
                gr.font.name = 'Arial'
            doc.add_paragraph()

    # ── Projects ───────────────────────────────────────────────────────
    if d.get('projects'):
        _section_header(doc, 'Projects')
        for proj in d['projects']:
            pp = doc.add_paragraph()
            pr = pp.add_run(proj.get('name', ''))
            pr.bold = True
            pr.font.size = Pt(10)
            pr.font.name = 'Arial'
            if proj.get('link'):
                pl = pp.add_run(f"  ({proj['link']})")
                pl.font.size = Pt(9)
                pl.font.color.rgb = ACCENT
                pl.font.name = 'Arial'
            if proj.get('description'):
                dp = doc.add_paragraph()
                dr = dp.add_run(proj['description'])
                dr.font.size = Pt(10)
                dr.font.name = 'Arial'
            for b in proj.get('bullets', []):
                _bullet(doc, b)
            doc.add_paragraph()

    # ── Certifications ─────────────────────────────────────────────────
    if d.get('certifications'):
        _section_header(doc, 'Certifications')
        for cert in d['certifications']:
            cp = doc.add_paragraph()
            cp.paragraph_format.space_after = Pt(4)
            cr1 = cp.add_run(cert.get('name', ''))
            cr1.bold = True
            cr1.font.size = Pt(10)
            cr1.font.name = 'Arial'
            if cert.get('issuer'):
                cr2 = cp.add_run(f"  \u2014  {cert['issuer']}")
                cr2.font.size = Pt(10)
                cr2.font.color.rgb = GRAY
                cr2.font.name = 'Arial'
            if cert.get('year'):
                cr3 = cp.add_run(f"  ({cert['year']})")
                cr3.font.size = Pt(9)
                cr3.font.color.rgb = GRAY
                cr3.font.name = 'Arial'

    # ── Languages ──────────────────────────────────────────────────────
    if d.get('languages'):
        _section_header(doc, 'Languages')
        lp = doc.add_paragraph()
        lr = lp.add_run('  \u2022  '.join(d['languages']))
        lr.font.size = Pt(10)
        lr.font.name = 'Arial'

    # ── Achievements ───────────────────────────────────────────────────
    if d.get('achievements'):
        _section_header(doc, 'Achievements & Awards')
        for a in d['achievements']:
            _bullet(doc, a)

    # Return bytes
    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ── Misc helpers ─────────────────────────────────────────────────────────────

# WITH this:
def fetch_yt_video(link):
    try:
        ydl_opts = {
            'quiet': True,
            'skip_download': True,
            'extract_flat': True,
            'no_warnings': True,
        }
        with YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(link, download=False)
            return info.get('title', 'Watch Video')
    except Exception:
        return 'Watch Video'


def get_table_download_link(df, filename, text):
    csv = df.to_csv(index=False)
    b64 = base64.b64encode(csv.encode()).decode()
    return f'<a href="data:file/csv;base64,{b64}" download="{filename}">{text}</a>'


def pdf_reader(file):
    resource_manager = PDFResourceManager()
    fake_file_handle = io.StringIO()
    converter = TextConverter(resource_manager, fake_file_handle, laparams=LAParams())
    page_interpreter = PDFPageInterpreter(resource_manager, converter)
    with open(file, 'rb') as fh:
        for page in PDFPage.get_pages(fh, caching=True, check_extractable=True):
            page_interpreter.process_page(page)
    text = fake_file_handle.getvalue()
    converter.close()
    fake_file_handle.close()
    return text


def show_pdf(file_path):
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    pdf_display = (
        f'<iframe src="data:application/pdf;base64,{base64_pdf}" '
        f'width="700" height="1000" type="application/pdf"></iframe>'
    )
    st.markdown(pdf_display, unsafe_allow_html=True)


def course_recommender(course_list):
    st.subheader("**Courses & Certificates🎓 Recommendations**")
    c = 0
    rec_course = []
    no_of_reco = st.slider('Choose Number of Course Recommendations:', 1, 10, 4)
    random.shuffle(course_list)
    for c_name, c_link in course_list:
        c += 1
        st.markdown(
            f"({c}) <a href='{c_link}' style='color: red; text-decoration: none;'>{c_name}</a>",
            unsafe_allow_html=True)
        rec_course.append(c_name)
        if c == no_of_reco:
            break
    return rec_course


# ── Database ─────────────────────────────────────────────────────────────────

connection = pymysql.connect(host='localhost', user='root', password='')
cursor = connection.cursor()


def insert_data(name, email, res_score, timestamp, no_of_pages, reco_field,
                cand_level, skills, recommended_skills, courses):
    DB_table_name = 'user_data'
    insert_sql = (
        "insert into " + DB_table_name +
        " values (0,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
    )
    rec_values = (
        name, email, str(res_score), timestamp, str(no_of_pages),
        reco_field, cand_level, skills, recommended_skills, courses,
    )
    cursor.execute(insert_sql, rec_values)
    connection.commit()


# ── n8n Email Automation ──────────────────────────────────────────────────────

N8N_WEBHOOK_URL = "https://vedantkumar03.app.n8n.cloud/webhook/76e5a815-42c0-40bc-bba1-4dc9b9557af5"
# ↑ Replace with your actual n8n webhook URL after setting up the workflow

def send_report_via_n8n(recipient_email: str, report_data: dict) -> bool:
    """
    POST the analysis report payload to n8n webhook.
    n8n will format and send the email to the user.
    Returns True on success, False on failure.
    """
    import requests

    payload = {
        "recipient_email":    recipient_email,
        "candidate_name":     report_data.get("name", ""),
        "resume_score":       report_data.get("resume_score", 0),
        "candidate_level":    report_data.get("cand_level", ""),
        "predicted_field":    report_data.get("reco_field", ""),
        "current_skills":     report_data.get("skills", []),
        "recommended_skills": report_data.get("recommended_skills", []),
        "recommended_courses":report_data.get("rec_course", []),
        "timestamp":          report_data.get("timestamp", ""),
    }

    try:
        response = requests.post(
            N8N_WEBHOOK_URL,
            json=payload,
            timeout=10
        )
        return response.status_code == 200
    except requests.exceptions.ConnectionError:
        return False
    except requests.exceptions.Timeout:
        return False
    except Exception:
        return False


# ── Page config & CSS ─────────────────────────────────────────────────────────

st.set_page_config(page_title="HIRE READY", page_icon='./Logo/Logo.ico', layout="wide")

# ── Theme definitions ─────────────────────────────────────────────────────────

LIGHT_CSS = """
:root {
  --bg:        #FFFDF2;
  --bg2:       #FFFFFF;
  --text:      #111111;
  --text-muted:#555555;
  --accent:    #F5C518;
  --accent2:   #6B4FBB;
  --border:    #111111;
  --card-bg:   #FFFFFF;
  --input-bg:  #FFFFFF;
  --dot:       rgba(0,0,0,0.15);
  --shadow:    4px 4px 0 #111111;
  --btn-bg:    #111111;
  --btn-text:  #FFFFFF;
  --header-bg: #FFFDF2;
}
"""

DARK_CSS = """
:root {
  --bg:        #0E1117;
  --bg2:       #1A1D27;
  --text:      #F0F0F0;
  --text-muted:#AAAAAA;
  --accent:    #F5C518;
  --accent2:   #9B7FE8;
  --border:    #333333;
  --card-bg:   #1A1D27;
  --input-bg:  #262B38;
  --dot:       rgba(255,255,255,0.04);
  --shadow:    4px 4px 0 #F5C518;
  --btn-bg:    #F5C518;
  --btn-text:  #111111;
  --header-bg: #0E1117;
}
"""

SHARED_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Bebas+Neue&family=DM+Sans:wght@400;500;600&display=swap');

body,
[data-testid="stAppViewContainer"],
[data-testid="stApp"],
.stApp,
[data-testid="stMain"],
[data-testid="stMainBlockContainer"],
[data-testid="stBottomBlockContainer"],
.main, .block-container {
  background-color: var(--bg) !important;
  background-image: radial-gradient(var(--dot) 1.2px, transparent 1.2px) !important;
  background-size: 22px 22px !important;
  color: var(--text) !important;
  font-family: 'DM Sans', sans-serif !important;
}

[data-testid="stToolbar"],
[data-testid="stHeader"],
header[data-testid="stHeader"] {
  background-color: var(--header-bg) !important;
}

h1, h2, h3, h4, [data-testid="stMarkdownContainer"] h1,
[data-testid="stMarkdownContainer"] h2,
[data-testid="stMarkdownContainer"] h3 {
  color: var(--text) !important;
  font-family: 'DM Sans', sans-serif !important;
}

p, label, .stMarkdown p,
[data-testid="stMarkdownContainer"] p,
[data-testid="stWidgetLabel"] p { color: var(--text) !important; }

/* Buttons */
.stButton > button,
[data-testid="stBaseButton-secondary"],
[data-testid="stBaseButton-primary"] {
  background-color: var(--btn-bg) !important;
  color: var(--btn-text) !important;
  border: 2px solid var(--border) !important;
  border-radius: 8px !important;
  font-weight: 600 !important;
  box-shadow: 3px 3px 0 var(--accent) !important;
  transition: all 0.2s ease !important;
}
.stButton > button:hover {
  transform: translate(-2px, -2px) !important;
  box-shadow: 5px 5px 0 var(--accent) !important;
}

/* Inputs */
input, textarea,
[data-baseweb="input"] input,
[data-baseweb="textarea"] textarea {
  background-color: var(--input-bg) !important;
  color: var(--text) !important;
  border: 2px solid var(--border) !important;
  border-radius: 8px !important;
}
input::placeholder, textarea::placeholder { color: var(--text-muted) !important; }

/* Selectbox */
[data-testid="stSelectbox"] > div > div,
[data-baseweb="select"] > div {
  background-color: var(--input-bg) !important;
  color: var(--text) !important;
  border: 2px solid var(--border) !important;
  border-radius: 8px !important;
}
[data-baseweb="select"] span,
[data-baseweb="select"] div { color: var(--text) !important; }
[data-baseweb="popover"] ul,
[data-baseweb="menu"] { background-color: var(--card-bg) !important; border: 2px solid var(--border) !important; border-radius: 8px !important; }
[data-baseweb="menu"] li { color: var(--text) !important; }
[data-baseweb="menu"] li:hover { background-color: var(--accent) !important; color: #111 !important; }

/* File uploader */
[data-testid="stFileUploader"],
[data-testid="stFileUploaderDropzone"] {
  background-color: var(--card-bg) !important;
  border: 2.5px dashed var(--border) !important;
  border-radius: 16px !important;
}
[data-testid="stFileUploaderDropzoneInstructions"] span,
[data-testid="stFileUploaderDropzone"] span,
[data-testid="stFileUploaderDropzone"] small,
[data-testid="stFileUploaderDropzone"] p { color: var(--text) !important; }
[data-testid="stFileUploaderDropzone"] button { background-color: var(--btn-bg) !important; color: var(--btn-text) !important; border-radius: 8px !important; box-shadow: 3px 3px 0 var(--accent) !important; }

/* Number input */
[data-testid="stNumberInput"] input { background-color: var(--input-bg) !important; color: var(--text) !important; }

/* Expander */
[data-testid="stExpander"] {
  background-color: var(--card-bg) !important;
  border: 1px solid var(--border) !important;
  border-radius: 8px !important;
}
[data-testid="stExpander"] summary,
[data-testid="stExpander"] p { color: var(--text) !important; }

/* Tabs */
[data-testid="stTabs"] button[role="tab"] {
  color: var(--text) !important;
  background: transparent !important;
  border: none !important;
  box-shadow: none !important;
  font-weight: 600 !important;
}
[data-testid="stTabs"] button[role="tab"][aria-selected="true"] {
  border-bottom: 3px solid var(--accent) !important;
}

/* Divider */
hr { border-color: var(--border) !important; opacity: 0.3; }

/* Sidebar */
[data-testid="stSidebar"] {
  background-color: var(--bg2) !important;
  border-right: 1px solid var(--border) !important;
}

/* Form */
[data-testid="stForm"] {
  background-color: var(--card-bg) !important;
  border: 1px solid var(--border) !important;
  border-radius: 12px !important;
  padding: 1rem !important;
}

/* Number input buttons */
[data-testid="stNumberInput"] button {
  background-color: var(--btn-bg) !important;
  color: var(--btn-text) !important;
  border: 1px solid var(--border) !important;
}

/* Success / Error / Warning */
[data-testid="stAlert"] { border-radius: 8px !important; }

/* Theme toggle button special style */
.theme-toggle button {
  background: transparent !important;
  border: 2px solid var(--accent) !important;
  color: var(--text) !important;
  box-shadow: none !important;
  font-size: 1.2rem !important;
  padding: 0.3rem 0.8rem !important;
}
.theme-toggle button:hover {
  background: var(--accent) !important;
  color: #111 !important;
  transform: none !important;
  box-shadow: none !important;
}

/* Chat interface */
[data-testid="stChatMessage"] {
  background-color: var(--card-bg) !important;
  border: 1px solid var(--border) !important;
  border-radius: 12px !important;
  margin-bottom: 0.5rem !important;
  color: var(--text) !important;
}
[data-testid="stChatMessage"] p,
[data-testid="stChatMessage"] li,
[data-testid="stChatMessage"] span {
  color: var(--text) !important;
}
[data-testid="stChatInput"] {
  background-color: var(--input-bg) !important;
  border-color: var(--border) !important;
}
[data-testid="stChatInput"] textarea {
  color: var(--text) !important;
  background-color: var(--input-bg) !important;
}
"""

def apply_theme():
    """Inject the current theme CSS into the page."""
    if 'dark_mode' not in st.session_state:
        st.session_state.dark_mode = False
    theme_vars = DARK_CSS if st.session_state.dark_mode else LIGHT_CSS
    st.markdown(f"<style>{theme_vars}{SHARED_CSS}</style>", unsafe_allow_html=True)

def theme_toggle_button():
    """Render a small toggle in the top-right corner."""
    apply_theme()
    icon  = "☀️ Light" if st.session_state.dark_mode else "🌙 Dark"
    col_spacer, col_btn = st.columns([10, 1])
    with col_btn:
        st.markdown('<div class="theme-toggle">', unsafe_allow_html=True)
        if st.button(icon, key="theme_toggle"):
            st.session_state.dark_mode = not st.session_state.dark_mode
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)


# ── Generate Resume Tab ───────────────────────────────────────────────────────

def resume_generator_section():
    st.header("📝 Generate Your Resume")
    st.markdown("Fill in your details below and download a professionally formatted `.docx` resume in seconds.")

    with st.form("resume_form"):

        # Personal Info
        st.subheader("👤 Personal Information")
        col1, col2 = st.columns(2)
        with col1:
            name     = st.text_input("Full Name *",    placeholder="e.g. Vedant Sharma")
            email    = st.text_input("Email *",        placeholder="vedant@email.com")
            phone    = st.text_input("Phone",          placeholder="+91 98765 43210")
        with col2:
            linkedin = st.text_input("LinkedIn URL",   placeholder="linkedin.com/in/vedant")
            location = st.text_input("Location",       placeholder="Delhi, India")

        summary = st.text_area("Professional Summary",
                               placeholder="Brief 2-3 sentence overview of your profile and career goals.",
                               height=100)

        # Skills
        st.subheader("🛠️ Skills")
        skills_raw = st.text_input("Skills (comma-separated)",
                                   placeholder="Python, React, Machine Learning, SQL, Docker")

        # Work Experience
        st.subheader("💼 Work Experience")
        num_exp = st.number_input("Number of work experiences", min_value=0, max_value=10, value=3, step=1)
        experience = []
        for i in range(int(num_exp)):
            st.markdown(f"**Experience {i+1}**")
            c1, c2, c3 = st.columns(3)
            with c1:
                title     = st.text_input("Job Title",  key=f"exp_title_{i}",   placeholder="Software Engineer")
                company   = st.text_input("Company",    key=f"exp_company_{i}", placeholder="Tech Corp")
            with c2:
                exp_loc   = st.text_input("Location",   key=f"exp_loc_{i}",     placeholder="Delhi, India")
                exp_start = st.text_input("Start Date", key=f"exp_start_{i}",   placeholder="Jan 2022")
            with c3:
                exp_end   = st.text_input("End Date (blank = Present)", key=f"exp_end_{i}", placeholder="Present")
            exp_bullets_raw = st.text_area("Responsibilities / Achievements (one per line)",
                                           key=f"exp_bullets_{i}", height=100,
                                           placeholder="Built a REST API serving 10k+ requests/day.\nLed a team of 3 engineers.")
            exp_bullets = [b.strip() for b in exp_bullets_raw.splitlines() if b.strip()]
            if title or company:
                experience.append({
                    "title": title, "company": company, "location": exp_loc,
                    "start": exp_start, "end": exp_end or "Present",
                    "bullets": exp_bullets
                })
            st.divider()

        # Education
        st.subheader("🎓 Education")
        num_edu = st.number_input("Number of education entries", min_value=0, max_value=6, value=3, step=1)
        education = []
        for i in range(int(num_edu)):
            st.markdown(f"**Education {i+1}**")
            c1, c2, c3 = st.columns(3)
            with c1:
                degree      = st.text_input("Degree",      key=f"edu_deg_{i}",   placeholder="B.Tech Computer Science")
                institution = st.text_input("Institution", key=f"edu_inst_{i}",  placeholder="Delhi Technological University")
            with c2:
                edu_start   = st.text_input("Start Year",  key=f"edu_start_{i}", placeholder="2018")
                edu_end     = st.text_input("End Year",    key=f"edu_end_{i}",   placeholder="2022")
            with c3:
                gpa         = st.text_input("GPA / % (optional)", key=f"edu_gpa_{i}", placeholder="8.5 / 10")
            if degree or institution:
                education.append({
                    "degree": degree, "institution": institution,
                    "start": edu_start, "end": edu_end, "gpa": gpa
                })
            st.divider()

        # Projects
        st.subheader("🚀 Projects")
        num_proj = st.number_input("Number of projects", min_value=0, max_value=10, value=3, step=1)
        projects = []
        for i in range(int(num_proj)):
            st.markdown(f"**Project {i+1}**")
            c1, c2 = st.columns(2)
            with c1:
                proj_name = st.text_input("Project Name",               key=f"proj_name_{i}", placeholder="Hire Ready – Resume Analyzer")
                proj_link = st.text_input("GitHub / Live Link (opt.)",  key=f"proj_link_{i}", placeholder="github.com/you/project")
            with c2:
                proj_desc = st.text_input("One-line Description",       key=f"proj_desc_{i}",
                                          placeholder="AI-powered resume analysis tool built with Streamlit.")
            proj_bullets_raw = st.text_area("Highlights (one per line)", key=f"proj_bullets_{i}", height=80,
                                            placeholder="Integrated spaCy NER for extraction.\nAchieved 90% skill detection accuracy.")
            proj_bullets = [b.strip() for b in proj_bullets_raw.splitlines() if b.strip()]
            if proj_name:
                projects.append({
                    "name": proj_name, "link": proj_link,
                    "description": proj_desc, "bullets": proj_bullets
                })
            st.divider()

        # Certifications
        st.subheader("📜 Certifications")
        certs_raw = st.text_area("Certifications (one per line: Name | Issuer | Year)", height=80,
                                 placeholder="AWS Certified Developer | Amazon | 2023\nGoogle Data Analytics | Google | 2022")
        certifications = []
        for line in certs_raw.splitlines():
            parts = [p.strip() for p in line.split('|')]
            if parts[0]:
                certifications.append({
                    "name":   parts[0],
                    "issuer": parts[1] if len(parts) > 1 else "",
                    "year":   parts[2] if len(parts) > 2 else "",
                })

        # Languages & Achievements
        st.subheader("🌐 Languages & 🏅 Achievements")
        c1, c2 = st.columns(2)
        with c1:
            languages_raw    = st.text_input("Languages (comma-separated)",
                                             placeholder="English (Fluent), Hindi (Native)")
        with c2:
            achievements_raw = st.text_area("Achievements (one per line)", height=80,
                                            placeholder="Winner – Smart India Hackathon 2021\nPublished paper at ICML 2022")

        submitted = st.form_submit_button("✨ Generate My Resume")

    if submitted:
        if not name or not email:
            st.error("Please fill in at least your Full Name and Email.")
            return

        resume_dict = {
            "name":           name,
            "email":          email,
            "phone":          phone,
            "linkedin":       linkedin,
            "location":       location,
            "summary":        summary,
            "skills":         [s.strip() for s in skills_raw.split(',') if s.strip()],
            "experience":     experience,
            "education":      education,
            "projects":       projects,
            "certifications": certifications,
            "languages":      [l.strip() for l in languages_raw.split(',') if l.strip()],
            "achievements":   [a.strip() for a in achievements_raw.splitlines() if a.strip()],
        }

        with st.spinner("Generating your resume..."):
            try:
                docx_bytes = generate_resume_docx(resume_dict)
                st.success("✅ Your resume is ready!")
                st.download_button(
                    label="⬇️ Download Resume (.docx)",
                    data=docx_bytes,
                    file_name=f"{name.replace(' ', '_')}_Resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.info("💡 Open in Microsoft Word or Google Docs to further customize it.")
            except Exception as e:
                st.error(f"Failed to generate resume: {e}")


# ── Main app ──────────────────────────────────────────────────────────────────

def run():
    theme_toggle_button()
    st.title("HIRE READY")

    # ── Pre-set Gemini API key (hardcoded — no user input needed) ────
    if 'gemini_api_key' not in st.session_state:
        st.session_state.gemini_api_key = "AIzaSyCPUgqMGU8qA8MTgPGm1XnrSC5ic5iWaGQ"

    # ── Sidebar: show AI assistant status ───────────────────────────
    with st.sidebar:
        st.markdown("### 🤖 AI Career Assistant")
        st.success("✅ AI Assistant is active and ready!")
        st.markdown(
            "Switch to the **🤖 AI Career Assistant** tab to chat with your "
            "personal career advisor powered by Google Gemini."
        )
        st.markdown("---")

    with st.expander("Choose User", expanded=False):
        activities = ["Normal User", "Admin"]
        choice = st.selectbox("Choose among the given options:", activities)

    img = Image.open('./Logo/Logo.png')
    img = img.resize((2560, 1440))
    st.image(img)

    cursor.execute("CREATE DATABASE IF NOT EXISTS data_2;")
    connection.select_db("data_2")
    DB_table_name = 'user_data'
    table_sql = (
        "CREATE TABLE IF NOT EXISTS " + DB_table_name + """
        (ID INT NOT NULL AUTO_INCREMENT,
         Name varchar(100) NOT NULL,
         Email_ID VARCHAR(50) NOT NULL,
         resume_score VARCHAR(8) NOT NULL,
         Timestamp VARCHAR(50) NOT NULL,
         Page_no VARCHAR(5) NOT NULL,
         Predicted_Field VARCHAR(25) NOT NULL,
         User_level VARCHAR(30) NOT NULL,
         Actual_skills VARCHAR(300) NOT NULL,
         Recommended_skills VARCHAR(300) NOT NULL,
         Recommended_courses VARCHAR(600) NOT NULL,
         PRIMARY KEY (ID));"""
    )
    cursor.execute(table_sql)

    if choice == 'Normal User':

        # Three tabs: Analyze, Generate, AI Chat
        tab1, tab2, tab3 = st.tabs(["📊 Analyze My Resume", "✨ Generate New Resume", "🤖 AI Career Assistant"])

        with tab1:
            st.subheader("Upload & Analyze Your Resume")
            pdf_file = st.file_uploader("Choose your Resume", type=["pdf"])
            if pdf_file is not None:
                save_image_path = './Uploaded_Resumes/' + pdf_file.name
                with open(save_image_path, "wb") as f:
                    f.write(pdf_file.getbuffer())
                show_pdf(save_image_path)

                resume_data = parse_resume(save_image_path)

                if resume_data:
                    resume_text = pdf_reader(save_image_path)
                    # Store resume context for the chatbot tab
                    st.session_state.resume_data = resume_data
                    st.session_state.resume_text = resume_text
                    st.header("**Resume Analysis**")
                    st.success("Hello " + resume_data['name'])
                    st.subheader("**Your Basic info**")
                    try:
                        st.markdown(f"**Name:** {resume_data['name']}")
                        st.markdown(f"**Email:** {resume_data['email']}")
                        st.markdown(f"**Contact:** {resume_data['mobile_number']}")
                        st.markdown(f"**Resume pages:** {str(resume_data['no_of_pages'])}")
                    except Exception:
                        pass

                    cand_level = ''
                    if resume_data['no_of_pages'] == 1:
                        cand_level = "Fresher"
                        st.markdown('''<h4 style='text-align: left; color: #d73b5c;'>You are looking Fresher.</h4>''', unsafe_allow_html=True)
                    elif resume_data['no_of_pages'] == 2:
                        cand_level = "Intermediate"
                        st.markdown('''<h4 style='text-align: left; color: #1ed760;'>You are at intermediate level!</h4>''', unsafe_allow_html=True)
                    elif resume_data['no_of_pages'] >= 3:
                        cand_level = "Experienced"
                        st.markdown('''<h4 style='text-align: left; color: #fba171;'>You are at experience level!</h4>''', unsafe_allow_html=True)

                    st.subheader("**Skills Recommendation💡**")
                    keywords = st_tags(label='### Skills that you have',
                                       text='See our skills recommendation',
                                       value=resume_data['skills'], key='1')

                    ds_keyword      = ['tensorflow', 'keras', 'pytorch', 'machine learning', 'deep learning', 'flask', 'streamlit']
                    web_keyword     = ['react', 'django', 'node js', 'node.js', 'react js', 'php', 'laravel', 'magento', 'wordpress', 'javascript', 'angular js', 'c#', 'flask']
                    android_keyword = ['android', 'android development', 'flutter', 'kotlin', 'xml', 'kivy']
                    ios_keyword     = ['ios', 'ios development', 'swift', 'cocoa', 'cocoa touch', 'xcode']
                    uiux_keyword    = ['ux', 'adobe xd', 'figma', 'zeplin', 'balsamiq', 'ui', 'prototyping', 'wireframes',
                                       'storyframes', 'adobe photoshop', 'photoshop', 'editing', 'adobe illustrator',
                                       'illustrator', 'adobe after effects', 'after effects', 'adobe premier pro',
                                       'premier pro', 'adobe indesign', 'indesign', 'wireframe', 'solid', 'grasp',
                                       'user research', 'user experience']

                    recommended_skills = []
                    reco_field = ''
                    rec_course = ''

                    for i in resume_data['skills']:
                        il = i.lower()
                        if il in ds_keyword:
                            reco_field = 'Data Science'
                            st.success("** Our analysis says you are looking for Data Science Jobs.**")
                            recommended_skills = ['Data Visualization', 'Predictive Analysis', 'Statistical Modeling',
                                                  'Data Mining', 'Clustering & Classification', 'Data Analytics',
                                                  'Quantitative Analysis', 'Web Scraping', 'ML Algorithms', 'Keras',
                                                  'Pytorch', 'Probability', 'Scikit-learn', 'Tensorflow', 'Flask', 'Streamlit']
                            st_tags(label='### Recommended skills for you.', text='Recommended skills generated from System',
                                    value=recommended_skills, key='2')
                            st.markdown('''<h4 style='text-align: left; color: #1ed760;'>Adding this skills to resume will boost🚀 the chances of getting a Job💼</h4>''', unsafe_allow_html=True)
                            rec_course = course_recommender(ds_course)
                            break
                        elif il in web_keyword:
                            reco_field = 'Web Development'
                            st.success("** Our analysis says you are looking for Web Development Jobs **")
                            recommended_skills = ['React', 'Django', 'Node JS', 'React JS', 'php', 'laravel', 'Magento', 'wordpress', 'Javascript', 'Angular JS', 'c#', 'Flask', 'SDK']
                            st_tags(label='### Recommended skills for you.', text='Recommended skills generated from System',
                                    value=recommended_skills, key='3')
                            st.markdown('''<h4 style='text-align: left; color: #1ed760;'>Adding this skills to resume will boost🚀 the chances of getting a Job💼</h4>''', unsafe_allow_html=True)
                            rec_course = course_recommender(web_course)
                            break
                        elif il in android_keyword:
                            reco_field = 'Android Development'
                            st.success("** Our analysis says you are looking for Android App Development Jobs **")
                            recommended_skills = ['Android', 'Android development', 'Flutter', 'Kotlin', 'XML', 'Java', 'Kivy', 'GIT', 'SDK', 'SQLite']
                            st_tags(label='### Recommended skills for you.', text='Recommended skills generated from System',
                                    value=recommended_skills, key='4')
                            st.markdown('''<h4 style='text-align: left; color: #1ed760;'>Adding this skills to resume will boost🚀 the chances of getting a Job💼</h4>''', unsafe_allow_html=True)
                            rec_course = course_recommender(android_course)
                            break
                        elif il in ios_keyword:
                            reco_field = 'IOS Development'
                            st.success("** Our analysis says you are looking for IOS App Development Jobs **")
                            recommended_skills = ['IOS', 'IOS Development', 'Swift', 'Cocoa', 'Cocoa Touch', 'Xcode', 'Objective-C', 'SQLite', 'Plist', 'StoreKit', 'UI-Kit', 'AV Foundation', 'Auto-Layout']
                            st_tags(label='### Recommended skills for you.', text='Recommended skills generated from System',
                                    value=recommended_skills, key='5')
                            st.markdown('''<h4 style='text-align: left; color: #1ed760;'>Adding this skills to resume will boost🚀 the chances of getting a Job💼</h4>''', unsafe_allow_html=True)
                            rec_course = course_recommender(ios_course)
                            break
                        elif il in uiux_keyword:
                            reco_field = 'UI-UX Development'
                            st.success("** Our analysis says you are looking for UI-UX Development Jobs **")
                            recommended_skills = ['UI', 'User Experience', 'Adobe XD', 'Figma', 'Zeplin', 'Balsamiq', 'Prototyping', 'Wireframes', 'Storyframes', 'Adobe Photoshop', 'Editing', 'Illustrator', 'After Effects', 'Premier Pro', 'Indesign', 'Wireframe', 'Solid', 'Grasp', 'User Research']
                            st_tags(label='### Recommended skills for you.', text='Recommended skills generated from System',
                                    value=recommended_skills, key='6')
                            st.markdown('''<h4 style='text-align: left; color: #1ed760;'>Adding this skills to resume will boost🚀 the chances of getting a Job💼</h4>''', unsafe_allow_html=True)
                            rec_course = course_recommender(uiux_course)
                            break

                    ts = time.time()
                    cur_date  = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d')
                    cur_time  = datetime.datetime.fromtimestamp(ts).strftime('%H:%M:%S')
                    timestamp = cur_date + '_' + cur_time

                    st.subheader("**Resume Tips & Ideas💡**")
                    resume_score = 0
                    if 'Objective' in resume_text:
                        resume_score += 20
                        st.markdown('''<h4 style='text-align: left; color: #1ed760;'>[+] Awesome! You have added Objective</h4>''', unsafe_allow_html=True)
                    else:
                        st.markdown('''<h4 style='text-align: left; color: #fabc10;'>[-] Please add your career objective.</h4>''', unsafe_allow_html=True)
                    if 'Declaration' in resume_text:
                        resume_score += 20
                        st.markdown('''<h4 style='text-align: left; color: #1ed760;'>[+] Awesome! You have added Declaration✍</h4>''', unsafe_allow_html=True)
                    else:
                        st.markdown('''<h4 style='text-align: left; color: #fabc10;'>[-] Please add a Declaration✍.</h4>''', unsafe_allow_html=True)
                    if 'Hobbies' in resume_text or 'Interests' in resume_text:
                        resume_score += 20
                        st.markdown('''<h4 style='text-align: left; color: #1ed760;'>[+] Awesome! You have added your Hobbies⚽</h4>''', unsafe_allow_html=True)
                    else:
                        st.markdown('''<h4 style='text-align: left; color: #fabc10;'>[-] Please add Hobbies⚽.</h4>''', unsafe_allow_html=True)
                    if 'Achievements' in resume_text:
                        resume_score += 20
                        st.markdown('''<h4 style='text-align: left; color: #1ed760;'>[+] Awesome! You have added your Achievements🏅</h4>''', unsafe_allow_html=True)
                    else:
                        st.markdown('''<h4 style='text-align: left; color: #fabc10;'>[-] Please add Achievements🏅.</h4>''', unsafe_allow_html=True)
                    if 'Projects' in resume_text:
                        resume_score += 20
                        st.markdown('''<h4 style='text-align: left; color: #1ed760;'>[+] Awesome! You have added your Projects👨‍💻</h4>''', unsafe_allow_html=True)
                    else:
                        st.markdown('''<h4 style='text-align: left; color: #fabc10;'>[-] Please add Projects👨‍💻.</h4>''', unsafe_allow_html=True)

                    st.subheader("**Resume Score📝**")
                    st.markdown("""<style>.stProgress > div > div > div > div { background-color: #d73b5c; }</style>""", unsafe_allow_html=True)
                    my_bar = st.progress(0)
                    score = 0
                    for percent_complete in range(resume_score):
                        score += 1
                        time.sleep(0.1)
                        my_bar.progress(percent_complete + 1)
                    st.success('** Your Resume Writing Score: ' + str(score) + '**')
                    st.warning("** Note: This score is calculated based on the content in your Resume. **")

                    # Store analysis results for chatbot context
                    st.session_state.cand_level = cand_level
                    st.session_state.resume_score = resume_score
                    st.session_state.reco_field = reco_field

                    insert_data(resume_data['name'], resume_data['email'], str(resume_score), timestamp,
                                str(resume_data['no_of_pages']), reco_field, cand_level,
                                str(resume_data['skills']), str(recommended_skills), str(rec_course))

                    st.header("**Bonus Video for Interview👨‍💼 Tips💡**")
                    interview_vid = random.choice(interview_videos)
                    st.subheader("✅ **" + fetch_yt_video(interview_vid) + "**")
                    st.video(interview_vid)

                    st.header("**Bonus Video for Resume Writing Tips💡**")
                    resume_vid = random.choice(resume_videos)
                    st.subheader("✅ **" + fetch_yt_video(resume_vid) + "**")
                    st.video(resume_vid)

                    connection.commit()

                    st.header("**ATS-Friendly Resume Template**")
                    st.markdown('''<h4 style='text-align: left; color: #1ed760;'>Would you like to download an ATS-friendly resume template?</h4>''', unsafe_allow_html=True)
                    if st.button("Download ATS-Friendly Resume Template"):
                        resume_path = r"C:\Users\vedan\Hire Ready\ATS Friendly Resume\ATS Friendly CV Template.pdf"
                        try:
                            with open(resume_path, "rb") as pdf_file:
                                PDFbyte = pdf_file.read()
                            st.download_button(label="Click here to download", data=PDFbyte,
                                               file_name="ATS_Friendly_Resume_Template.pdf",
                                               mime="application/octet-stream")
                            st.success("Download started! Check your downloads folder.")
                        except Exception as e:
                            st.error(f"Error loading resume template: {str(e)}")

                    # ── Email Report via n8n ──────────────────────────────
                    st.divider()
                    st.header("**📧 Get Your Report by Email**")
                    st.markdown(
                        '''<h4 style='text-align: left; color: #1ed760;'>
                        Enter your email below to receive a full copy of your analysis report.
                        </h4>''',
                        unsafe_allow_html=True
                    )

                    email_input = st.text_input(
                        "Your Email Address",
                        placeholder="example@gmail.com",
                        key="report_email_input"
                    )

                    if st.button("📨 Send Report to My Email"):
                        if not email_input or "@" not in email_input:
                            st.error("Please enter a valid email address.")
                        else:
                            report_payload = {
                                "name":               resume_data.get('name', ''),
                                "resume_score":       resume_score,
                                "cand_level":         cand_level,
                                "reco_field":         reco_field,
                                "skills":             resume_data.get('skills', []),
                                "recommended_skills": recommended_skills,
                                "rec_course":         rec_course,
                                "timestamp":          timestamp,
                            }
                            with st.spinner("Sending your report..."):
                                success = send_report_via_n8n(email_input, report_payload)
                            if success:
                                st.success(f"✅ Report sent successfully to **{email_input}**! Check your inbox.")
                                st.balloons()
                            else:
                                st.error(
                                    "❌ Could not reach the email service. "
                                    "Please check your n8n webhook URL or try again later."
                                )
                else:
                    st.error('Something went wrong..')

        with tab2:
            resume_generator_section()

        with tab3:
            render_chatbot(
                resume_data=st.session_state.get('resume_data'),
                resume_text=st.session_state.get('resume_text'),
            )

    else:
        st.success('Welcome to Admin Side')
        ad_user     = st.text_input("Username")
        ad_password = st.text_input("Password", type='password')
        if st.button('Login'):
            if ad_user == 'admin' and ad_password == 'admin123':
                st.success("Welcome")
                cursor.execute('SELECT * FROM user_data')
                data = cursor.fetchall()
                st.header("**User's👨‍💻 Data**")
                df = pd.DataFrame(data, columns=['ID', 'Name', 'Email', 'Resume Score', 'Timestamp',
                                                  'Total Page', 'Predicted Field', 'User Level',
                                                  'Actual Skills', 'Recommended Skills', 'Recommended Course'])
                st.dataframe(df)
                st.markdown(get_table_download_link(df, 'User_Data.csv', 'Download Report'), unsafe_allow_html=True)

                plot_data = pd.read_sql('SELECT * FROM user_data;', connection)

                labels = plot_data.Predicted_Field.unique()
                values = plot_data.Predicted_Field.value_counts()
                st.subheader("📈 **Pie-Chart for Predicted Field Recommendations**")
                fig = px.pie(df, values=values, names=labels, title='Predicted Field according to the Skills')
                st.plotly_chart(fig)

                labels = plot_data.User_level.unique()
                values = plot_data.User_level.value_counts()
                st.subheader("📈 **Pie-Chart for User's👨‍💻 Experienced Level**")
                fig = px.pie(df, values=values, names=labels, title="Pie-Chart📈 for User's👨‍💻 Experienced Level")
                st.plotly_chart(fig)
            else:
                st.error("Wrong ID & Password Provided")


run()